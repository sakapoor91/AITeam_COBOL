"""
Converts a BIZ-*.md file to a styled .docx Word document.
Usage: python _md_to_docx.py <input.md> [output.docx]
"""

import sys
import re
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Colour palette ──────────────────────────────────────────────────────────
H1_COLOR   = RGBColor(0x1F, 0x49, 0x7D)   # dark navy
H2_COLOR   = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
H3_COLOR   = RGBColor(0x2E, 0x74, 0xB5)
TBL_HDR_BG = "2E74B5"                      # blue header row
TBL_ALT_BG = "D6E4F0"                      # light blue alt row
CODE_BG    = "F2F2F2"                      # light grey code block
BORDER_CLR = "BFBFBF"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def set_table_borders(table):
    tbl     = table._tbl
    tblPr   = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblBord = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BORDER_CLR)
        tblBord.append(el)
    tblPr.append(tblBord)

def parse_inline(text):
    """Return list of (text, bold, italic, code) run tuples."""
    result = []
    # split on inline code first
    parts = re.split(r'(`[^`\n]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`') and len(part) > 2:
            result.append((part[1:-1], False, False, True))
        else:
            # split on bold (**) and italic (*)
            subparts = re.split(r'(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*)', part)
            for sp in subparts:
                if sp.startswith('**') and sp.endswith('**') and len(sp) > 4:
                    result.append((sp[2:-2], True, False, False))
                elif sp.startswith('*') and sp.endswith('*') and len(sp) > 2:
                    result.append((sp[1:-1], False, True, False))
                elif sp:
                    result.append((sp, False, False, False))
    return result

def add_runs(para, text, base_bold=False, base_size=None, base_color=None,
             code_font='Courier New'):
    """Add formatted runs to a paragraph, respecting inline markdown."""
    for (chunk, bold, italic, code) in parse_inline(text):
        run = para.add_run(chunk)
        run.bold   = bold or base_bold
        run.italic = italic
        if code:
            run.font.name = code_font
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            run.font.name = 'Calibri'
            if base_size:
                run.font.size = base_size
            if base_color:
                run.font.color.rgb = base_color

def add_heading(doc, text, level):
    colors = {1: H1_COLOR, 2: H2_COLOR, 3: H3_COLOR}
    sizes  = {1: Pt(18),   2: Pt(14),   3: Pt(12)}
    para   = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size  = sizes.get(level, Pt(11))
    run.font.color.rgb = colors.get(level, H3_COLOR)
    run.font.name  = 'Calibri'
    if level == 1:
        # underline for H1
        run.underline = True
    return para

def add_mermaid_diagram(doc, mermaid_lines, diagram_png_path):
    """
    Render a Mermaid diagram to a standalone PNG file alongside the DOCX,
    then insert a styled reference callout into the document (no embedded image).
    """
    src     = '\n'.join(mermaid_lines)
    png_out = diagram_png_path          # e.g.  BIZ-CBACT01C-flow.png
    rendered_ok = False

    tmp_mmd = None
    try:
        fd, tmp_mmd = tempfile.mkstemp(suffix='.mmd')
        os.close(fd)
        with open(tmp_mmd, 'w', encoding='utf-8') as f:
            f.write(src)
        npm_global = os.path.join(os.environ.get('APPDATA', ''), 'npm', 'mmdc.cmd')
        mmdc_cmd   = npm_global if os.path.exists(npm_global) else 'mmdc'
        result = subprocess.run(
            [mmdc_cmd, '-i', tmp_mmd, '-o', png_out,
             '-w', '3200', '-b', 'white', '--scale', '2'],
            capture_output=True, text=True, timeout=90,
            shell=(os.name == 'nt')
        )
        if result.returncode == 0 and os.path.exists(png_out):
            rendered_ok = True
            print(f"  Mermaid diagram saved: {os.path.basename(png_out)}")
        else:
            print(f"  mmdc warning: {result.stderr.strip()}", file=sys.stderr)
    except Exception as exc:
        print(f"  Mermaid render error: {exc}", file=sys.stderr)
    finally:
        if tmp_mmd and os.path.exists(tmp_mmd):
            try:
                os.unlink(tmp_mmd)
            except OSError:
                pass

    # ── Insert a styled reference callout into the DOCX ──────────────────────
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.left_indent  = Cm(0.5)
    set_para_bg(para, 'EEF4FB')

    icon = para.add_run('◆  ')
    icon.bold = True
    icon.font.color.rgb = H2_COLOR
    icon.font.size = Pt(10)

    if rendered_ok:
        label = para.add_run(f'Process Flow Diagram — see ')
        label.font.size = Pt(10)
        fname = para.add_run(os.path.basename(png_out))
        fname.bold = True
        fname.font.name = 'Courier New'
        fname.font.size = Pt(10)
        tail = para.add_run(
            '  (saved alongside this document).\n'
            'Open the PNG for full-resolution viewing, or open the source .md file\n'
            'in VS Code / any Mermaid-compatible viewer to explore the interactive diagram.'
        )
        tail.font.size = Pt(9)
        tail.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    else:
        label = para.add_run(
            'Process Flow Diagram — Mermaid source below.\n'
            'Paste at mermaid.live, or open the source .md file in VS Code to view.'
        )
        label.font.size = Pt(10)
        label.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        doc.add_paragraph()
        add_code_block(doc, mermaid_lines)


def add_code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.5)
        set_para_bg(para, CODE_BG.replace('#', ''))
        run = para.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x24, 0x24, 0x24)
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(4)

def add_table(doc, rows):
    """rows = list of lists of str. First row is header; second is separator (skip)."""
    if len(rows) < 2:
        return
    # filter out separator rows (cells are all dashes/spaces)
    data_rows = [r for r in rows if not all(re.match(r'^[-:\s]+$', c) for c in r)]
    if not data_rows:
        return
    ncols = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    set_table_borders(table)

    for ri, row in enumerate(data_rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            text = row[ci].strip() if ci < len(row) else ''
            add_runs(para, text, base_size=Pt(9))
            if ri == 0:
                set_cell_bg(cell, TBL_HDR_BG)
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif ri % 2 == 0:
                set_cell_bg(cell, TBL_ALT_BG)

    doc.add_paragraph()  # spacing after table

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    indent = Cm(0.5 + level * 0.5)
    para.paragraph_format.left_indent    = indent
    para.paragraph_format.first_line_indent = Cm(-0.4)
    bullet_run = para.add_run('• ')
    bullet_run.font.name  = 'Calibri'
    bullet_run.font.size  = Pt(10)
    bullet_run.font.color.rgb = H2_COLOR
    add_runs(para, text, base_size=Pt(10))
    return para

def add_hrule(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), H2_COLOR.__str__().replace('(', '').replace(')', '').replace(', ', ''))
    pBdr.append(bot)
    pPr.append(pBdr)

def parse_table_row(line):
    """Parse a markdown table row into a list of cell strings."""
    line = line.strip().strip('|')
    return [c.strip() for c in line.split('|')]

# ── Main converter ───────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    # Derive companion PNG filename: BIZ-CBACT01C-flow.png
    stem       = os.path.splitext(docx_path)[0]
    flow_png   = stem + '-flow.png'

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # ── Heading ─────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line) and not line.startswith('|'):
            add_hrule(doc)
            i += 1
            continue

        # ── Fenced code block (plain or mermaid) ────────────────────────────
        if line.strip().startswith('```'):
            fence_tag  = line.strip()[3:].strip().lower()  # e.g. 'mermaid', 'python', ''
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            if fence_tag == 'mermaid':
                add_mermaid_diagram(doc, code_lines, flow_png)
            else:
                add_code_block(doc, code_lines)
            continue

        # ── Table ────────────────────────────────────────────────────────────
        if line.startswith('|'):
            table_rows = []
            while i < n and lines[i].startswith('|'):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=indent)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=indent)
            i += 1
            continue

        # ── Blockquote (>) ───────────────────────────────────────────────────
        m = re.match(r'^>\s*(.*)', line)
        if m:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            add_runs(para, m.group(1), base_size=Pt(10),
                     base_color=RGBColor(0x40, 0x40, 0x40))
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if not line.strip():
            # Only add a gap if the previous paragraph wasn't already a spacer
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(4)
        add_runs(para, line, base_size=Pt(10))
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python _md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    md   = sys.argv[1]
    docx = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(md)[0] + '.docx'
    convert(md, docx)
